#!/usr/bin/env python3
"""
Convert DEPLOYMENT_ENDPOINTS.md to Word document with proper formatting
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
import re
import os

def create_word_document():
    """Create a comprehensive Word document from markdown"""
    
    # Create document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15
    
    # Add title
    title = doc.add_heading('Deployment Endpoints Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('PRAN Chatbot AWS - Complete API Endpoints Reference')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].bold = True
    
    intro = doc.add_paragraph(
        'This document provides detailed information about all deployment endpoints, '
        'API URLs, and integration points for the PRAN Chatbot system deployed on AWS. '
        'This is a comprehensive guide for developers to integrate with the chatbot API.'
    )
    intro.paragraph_format.space_after = Pt(12)
    
    # Read markdown file
    md_file = 'docs/DEPLOYMENT_ENDPOINTS.md'
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    i = 0
    in_code_block = False
    code_block_lines = []
    code_language = ''
    in_table = False
    table_lines = []
    
    while i < len(lines):
        line = lines[i]
        original_line = line
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                if code_block_lines:
                    code_text = '\n'.join(code_block_lines)
                    add_code_paragraph(doc, code_text, code_language)
                code_block_lines = []
                code_language = ''
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
                code_language = line.strip()[3:].strip()
            i += 1
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        
        # Handle tables
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
            table_lines.append(line)
            i += 1
            continue
        else:
            if in_table and table_lines:
                add_markdown_table(doc, table_lines)
                table_lines = []
                in_table = False
        
        # Skip horizontal rules (add spacing instead)
        if line.strip() == '---':
            doc.add_paragraph()
            i += 1
            continue
        
        # Handle headings
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            # Remove emoji and clean text
            text = re.sub(r'[📋🌐💻🔧🤖⚙️🏗️🔐🧪🔒📊🔄📝🐛📚📅]', '', text).strip()
            heading = doc.add_heading(text, level=min(level, 3))
            heading.style.font.name = 'Calibri'
            if level == 1:
                heading.style.font.size = Pt(16)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif level == 2:
                heading.style.font.size = Pt(14)
            else:
                heading.style.font.size = Pt(12)
            i += 1
            continue
        
        # Handle list items
        if re.match(r'^[\s]*[-*]\s', line):
            text = re.sub(r'^[\s]*[-*]\s', '', line).strip()
            text = clean_markdown(text)
            para = doc.add_paragraph(text, style='List Bullet')
            para.paragraph_format.left_indent = Inches(0.25)
            i += 1
            continue
        
        # Handle numbered list
        if re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            text = clean_markdown(text)
            para = doc.add_paragraph(text, style='List Number')
            para.paragraph_format.left_indent = Inches(0.25)
            i += 1
            continue
        
        # Regular paragraph
        if line.strip():
            text = clean_markdown(line)
            if text.strip():
                para = doc.add_paragraph(text)
        else:
            # Empty line - add spacing
            doc.add_paragraph()
        
        i += 1
    
    # Add footer with document info
    doc.add_page_break()
    footer_para = doc.add_paragraph('Document Information')
    footer_para.runs[0].font.bold = True
    footer_para.runs[0].font.size = Pt(12)
    
    doc.add_paragraph('Last Updated: 2024-12-02', style='List Bullet')
    doc.add_paragraph('Version: 1.0.0', style='List Bullet')
    doc.add_paragraph('Repository: https://github.com/PranDotAI1/Pran_bot_aws.git', style='List Bullet')
    
    # Save document
    output_file = 'docs/DEPLOYMENT_ENDPOINTS.docx'
    doc.save(output_file)
    
    file_size = os.path.getsize(output_file) / 1024
    print(f"✅ Word document created successfully!")
    print(f"📄 File: {output_file}")
    print(f"📊 Size: {file_size:.2f} KB")
    print(f"📝 Ready for developer use!")

def clean_markdown(text):
    """Clean markdown formatting from text"""
    # Remove bold
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    # Remove inline code (keep the text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    # Remove links (keep text)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    # Remove emojis
    text = re.sub(r'[📋🌐💻🔧🤖⚙️🏗️🔐🧪🔒📊🔄📝🐛📚📅✅❌⚠️]', '', text)
    return text.strip()

def add_code_paragraph(doc, code_text, language=''):
    """Add a code block with proper formatting"""
    para = doc.add_paragraph()
    para.style = 'No Spacing'
    para_format = para.paragraph_format
    para_format.left_indent = Inches(0.5)
    para_format.right_indent = Inches(0.5)
    para_format.space_before = Pt(6)
    para_format.space_after = Pt(6)
    
    # Add code text
    run = para.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Add border (simulated with shading)
    shading_elm = para._element.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F5F5F5')
    shading_elm.append(shading)

def add_markdown_table(doc, table_lines):
    """Add a table from markdown table format"""
    if len(table_lines) < 2:
        return
    
    # Parse header
    header_line = table_lines[0]
    separator_line = table_lines[1]
    
    # Extract headers
    headers = [h.strip() for h in header_line.split('|')[1:-1]]
    num_cols = len(headers)
    
    if num_cols == 0:
        return
    
    # Create table
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = 'Light Grid Accent 1'
    
    # Add header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = clean_markdown(header)
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    
    # Add data rows
    for line in table_lines[2:]:
        if '|' in line:
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if len(cells) == num_cols:
                row_cells = table.add_row().cells
                for i, cell_text in enumerate(cells):
                    row_cells[i].text = clean_markdown(cell_text)
                    row_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    
    # Add spacing after table
    doc.add_paragraph()

# Import needed for shading
from docx.oxml import OxmlElement

if __name__ == '__main__':
    create_word_document()

