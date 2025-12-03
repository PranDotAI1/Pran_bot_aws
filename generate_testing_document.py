#!/usr/bin/env python3
"""
Generate a comprehensive Word document for Pran.AI Chatbot Testing
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Style as hyperlink
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return hyperlink

def create_testing_document():
    """Create the comprehensive testing Word document"""
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('Pran.AI Chatbot Testing Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(0, 102, 204)
    title_run.font.size = Pt(24)
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('Comprehensive Test Cases & Prompts')
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()
    
    # Testing URL
    doc.add_heading('📍 Testing URL', 1)
    url_para = doc.add_paragraph('Amplify App: ')
    add_hyperlink(url_para, 'https://main.d1fw711o7cx5w2.amplifyapp.com/', 'https://main.d1fw711o7cx5w2.amplifyapp.com/')
    
    doc.add_paragraph()
    
    # Testing Objectives
    doc.add_heading('🎯 Testing Objectives', 1)
    objectives = [
        'Verify basic conversation flow and greetings',
        'Test doctor search and recommendations (RAG database integration)',
        'Validate appointment booking functionality',
        'Check insurance information retrieval',
        'Test symptom assessment capabilities',
        'Verify "Yes/No" handling (CRITICAL: prevent duplicate responses)',
        'Test all healthcare segments (Lab Results, Billing, Emergency, etc.)',
        'Validate error handling and edge cases'
    ]
    
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')
    
    doc.add_page_break()
    
    # Quick Test Prompts Section
    doc.add_heading('⚡ Quick Test Prompts (5 Minutes)', 1)
    
    quick_note = doc.add_paragraph()
    quick_note_run = quick_note.add_run('Copy and paste these prompts in order for rapid testing:')
    quick_note_run.font.bold = True
    quick_note_run.font.color.rgb = RGBColor(255, 0, 0)
    
    doc.add_paragraph()
    
    # Quick Test Cases
    quick_tests = [
        {
            'title': '1. Basic Flow',
            'tests': [
                {'prompt': 'Hello', 'expected': 'Greeting from Dr. AI ✅'},
                {'prompt': 'I am suffering from viral', 'expected': 'Suggests finding a doctor ✅'},
                {'prompt': 'yes', 'expected': 'Shows doctors list (SINGLE response, not 10!) ✅'}
            ]
        },
        {
            'title': '2. Doctor Search',
            'tests': [
                {'prompt': 'suggest me some doctors', 'expected': 'List of doctors from database ✅'},
                {'prompt': 'I need a gynecologist', 'expected': 'Shows gynecologists ✅'}
            ]
        },
        {
            'title': '3. Insurance',
            'tests': [
                {'prompt': 'all plans', 'expected': 'Shows all insurance plans (SINGLE response!) ✅'},
                {'prompt': 'tell me about insurance', 'expected': 'Insurance recommendations ✅'}
            ]
        },
        {
            'title': '4. Appointments',
            'tests': [
                {'prompt': 'I want to book an appointment', 'expected': 'Asks for details, shows doctors ✅'}
            ]
        },
        {
            'title': '5. Wellness',
            'tests': [
                {'prompt': 'I need diet recommendations', 'expected': 'Wellness guidance ✅'}
            ]
        },
        {
            'title': '6. Emergency',
            'tests': [
                {'prompt': 'This is an emergency', 'expected': 'Emergency guidance, 911 info ✅'}
            ]
        }
    ]
    
    for section in quick_tests:
        doc.add_heading(section['title'], 2)
        
        for test in section['tests']:
            # Prompt
            prompt_para = doc.add_paragraph()
            prompt_label = prompt_para.add_run('Prompt: ')
            prompt_label.font.bold = True
            prompt_label.font.color.rgb = RGBColor(0, 0, 255)
            prompt_text = prompt_para.add_run(f'"{test["prompt"]}"')
            prompt_text.font.italic = True
            
            # Expected
            expected_para = doc.add_paragraph()
            expected_label = expected_para.add_run('Expected: ')
            expected_label.font.bold = True
            expected_label.font.color.rgb = RGBColor(0, 128, 0)
            expected_text = expected_para.add_run(test['expected'])
            
            doc.add_paragraph()
    
    doc.add_page_break()
    
    # Detailed Test Cases
    doc.add_heading('🧪 Detailed Test Cases', 1)
    
    # Test Case 1: Basic Conversation & Greeting
    doc.add_heading('1. Basic Conversation & Greeting', 2)
    
    test_cases_basic = [
        {
            'id': '1.1',
            'name': 'Simple Greeting',
            'prompt': 'Hello',
            'expected': [
                'Single response (not duplicates)',
                'Greeting from Dr. AI',
                'Mentions capabilities (appointments, insurance, finding doctors, etc.)'
            ],
            'pass_criteria': 'Single response, friendly greeting'
        },
        {
            'id': '1.2',
            'name': 'Time-based Greeting',
            'prompt': 'Good morning',
            'expected': [
                'Appropriate greeting',
                'Offers help'
            ],
            'pass_criteria': 'Single response, contextual greeting'
        },
        {
            'id': '1.3',
            'name': 'How can you help?',
            'prompt': 'How can you help me?',
            'expected': [
                'List of capabilities',
                'Examples of what the bot can do',
                'Encouragement to ask questions'
            ],
            'pass_criteria': 'Comprehensive list of services'
        }
    ]
    
    for tc in test_cases_basic:
        doc.add_heading(f'Test Case {tc["id"]}: {tc["name"]}', 3)
        
        # Prompt
        p = doc.add_paragraph()
        p.add_run('Prompt: ').bold = True
        p.add_run(f'"{tc["prompt"]}"').italic = True
        
        # Expected Response
        p = doc.add_paragraph()
        p.add_run('Expected Response:').bold = True
        for exp in tc['expected']:
            doc.add_paragraph(exp, style='List Bullet 2')
        
        # Pass Criteria
        p = doc.add_paragraph()
        p.add_run('Pass Criteria: ').bold = True
        pass_run = p.add_run(f'✅ {tc["pass_criteria"]}')
        pass_run.font.color.rgb = RGBColor(0, 128, 0)
        
        doc.add_paragraph()
    
    # Test Case 2: Symptom Assessment & Doctor Recommendations
    doc.add_heading('2. Symptom Assessment & Doctor Recommendations', 2)
    
    test_cases_symptoms = [
        {
            'id': '2.1',
            'name': 'General Symptoms',
            'prompt': 'I am suffering from viral',
            'expected': [
                'Acknowledges health concern',
                'Offers to find appropriate doctor',
                'Asks if user wants to search for doctors'
            ],
            'pass_criteria': 'Empathetic response, actionable next step'
        },
        {
            'id': '2.2',
            'name': 'Affirmative Response (CRITICAL)',
            'setup': 'After Test Case 2.1',
            'prompt': 'yes',
            'expected': [
                '⚠️ SINGLE response only (NOT 10 duplicates)',
                'Shows list of available doctors',
                'Includes doctor details (name, specialty, contact)',
                'From database if available'
            ],
            'pass_criteria': 'SINGLE response, shows doctors list',
            'critical': True
        },
        {
            'id': '2.3',
            'name': 'Specific Symptom',
            'prompt': 'I have fever and cough',
            'expected': [
                'Recommends general physician',
                'Offers to find available doctors',
                'May show doctor list from database'
            ],
            'pass_criteria': 'Appropriate specialty recommendation'
        },
        {
            'id': '2.4',
            'name': 'Specific Specialty Request',
            'prompt': 'I need a gynecologist',
            'expected': [
                'Acknowledges request',
                'Shows available gynecologists from database',
                'Includes contact information',
                'Offers to book appointment'
            ],
            'pass_criteria': 'Shows gynecologist list from database'
        }
    ]
    
    for tc in test_cases_symptoms:
        heading = f'Test Case {tc["id"]}: {tc["name"]}'
        if tc.get('critical'):
            heading += ' ⚠️ CRITICAL'
        doc.add_heading(heading, 3)
        
        # Setup if applicable
        if 'setup' in tc:
            p = doc.add_paragraph()
            p.add_run('Setup: ').bold = True
            setup_run = p.add_run(tc['setup'])
            setup_run.font.color.rgb = RGBColor(255, 140, 0)
        
        # Prompt
        p = doc.add_paragraph()
        p.add_run('Prompt: ').bold = True
        p.add_run(f'"{tc["prompt"]}"').italic = True
        
        # Expected Response
        p = doc.add_paragraph()
        p.add_run('Expected Response:').bold = True
        for exp in tc['expected']:
            para = doc.add_paragraph(exp, style='List Bullet 2')
            if '⚠️' in exp or 'SINGLE' in exp or 'NOT' in exp:
                para.runs[0].font.color.rgb = RGBColor(255, 0, 0)
                para.runs[0].font.bold = True
        
        # Pass Criteria
        p = doc.add_paragraph()
        p.add_run('Pass Criteria: ').bold = True
        pass_run = p.add_run(f'✅ {tc["pass_criteria"]}')
        pass_run.font.color.rgb = RGBColor(0, 128, 0)
        if tc.get('critical'):
            pass_run.font.bold = True
        
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # Test Case 3: Doctor Search
    doc.add_heading('3. Doctor Search', 2)
    
    test_cases_doctors = [
        {
            'id': '3.1',
            'name': 'General Doctor Search',
            'prompt': 'suggest me some doctors',
            'expected': [
                'List of available doctors from database',
                'Doctor details (name, specialty, department, phone)',
                'Offers to book appointment'
            ],
            'pass_criteria': 'Shows multiple doctors with details'
        },
        {
            'id': '3.2',
            'name': 'Specialty-specific Search',
            'prompt': 'show me cardiologists',
            'expected': [
                'List of cardiologists from database',
                'Doctor details',
                'Booking option'
            ],
            'pass_criteria': 'Shows cardiologist-specific list'
        },
        {
            'id': '3.3',
            'name': 'All Doctors',
            'prompt': 'show all available doctors',
            'expected': [
                'Comprehensive list of doctors',
                'Multiple specialties',
                'Contact information'
            ],
            'pass_criteria': 'Shows extensive doctor list'
        }
    ]
    
    for tc in test_cases_doctors:
        doc.add_heading(f'Test Case {tc["id"]}: {tc["name"]}', 3)
        
        p = doc.add_paragraph()
        p.add_run('Prompt: ').bold = True
        p.add_run(f'"{tc["prompt"]}"').italic = True
        
        p = doc.add_paragraph()
        p.add_run('Expected Response:').bold = True
        for exp in tc['expected']:
            doc.add_paragraph(exp, style='List Bullet 2')
        
        p = doc.add_paragraph()
        p.add_run('Pass Criteria: ').bold = True
        pass_run = p.add_run(f'✅ {tc["pass_criteria"]}')
        pass_run.font.color.rgb = RGBColor(0, 128, 0)
        
        doc.add_paragraph()
    
    # Test Case 4: Insurance Information
    doc.add_heading('4. Insurance Information', 2)
    
    test_cases_insurance = [
        {
            'id': '4.1',
            'name': 'General Insurance Query',
            'prompt': 'I need help with insurance',
            'expected': [
                'Insurance plan recommendations',
                'Shows multiple plan options',
                'Details (premium, deductible, coverage)',
                'Asks if user wants more details'
            ],
            'pass_criteria': 'Shows insurance plans with details'
        },
        {
            'id': '4.2',
            'name': 'All Plans Request',
            'prompt': 'all plans',
            'expected': [
                'Comprehensive list of all insurance plans',
                'Basic Health Plan, Premium Health Plan, Family Health Plan',
                'Features and pricing',
                'Best suited for whom'
            ],
            'pass_criteria': 'Shows all plans with complete details',
            'critical': True
        },
        {
            'id': '4.3',
            'name': 'Plan Comparison',
            'prompt': 'compare insurance plans',
            'expected': [
                'Side-by-side comparison of plans',
                'Key differences highlighted',
                'Recommendations based on needs'
            ],
            'pass_criteria': 'Helpful comparison information'
        }
    ]
    
    for tc in test_cases_insurance:
        heading = f'Test Case {tc["id"]}: {tc["name"]}'
        if tc.get('critical'):
            heading += ' ⚠️ CRITICAL'
        doc.add_heading(heading, 3)
        
        p = doc.add_paragraph()
        p.add_run('Prompt: ').bold = True
        p.add_run(f'"{tc["prompt"]}"').italic = True
        
        p = doc.add_paragraph()
        p.add_run('Expected Response:').bold = True
        for exp in tc['expected']:
            doc.add_paragraph(exp, style='List Bullet 2')
        
        p = doc.add_paragraph()
        p.add_run('Pass Criteria: ').bold = True
        pass_run = p.add_run(f'✅ {tc["pass_criteria"]}')
        pass_run.font.color.rgb = RGBColor(0, 128, 0)
        
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # Test Case 5: Appointment Booking
    doc.add_heading('5. Appointment Booking', 2)
    
    test_cases_appointments = [
        {
            'id': '5.1',
            'name': 'Book Appointment Request',
            'prompt': 'I want to book an appointment',
            'expected': [
                'Asks for symptoms or specialty preference',
                'Offers to find available doctors',
                'Guides through booking process'
            ],
            'pass_criteria': 'Clear guidance for booking'
        },
        {
            'id': '5.2',
            'name': 'Appointment with Specific Doctor',
            'prompt': 'I want to book with Dr. Smith',
            'expected': [
                'Confirms doctor name',
                'Asks for preferred date/time',
                'Shows available slots (if integrated)'
            ],
            'pass_criteria': 'Confirms doctor, asks for details'
        }
    ]
    
    for tc in test_cases_appointments:
        doc.add_heading(f'Test Case {tc["id"]}: {tc["name"]}', 3)
        
        p = doc.add_paragraph()
        p.add_run('Prompt: ').bold = True
        p.add_run(f'"{tc["prompt"]}"').italic = True
        
        p = doc.add_paragraph()
        p.add_run('Expected Response:').bold = True
        for exp in tc['expected']:
            doc.add_paragraph(exp, style='List Bullet 2')
        
        p = doc.add_paragraph()
        p.add_run('Pass Criteria: ').bold = True
        pass_run = p.add_run(f'✅ {tc["pass_criteria"]}')
        pass_run.font.color.rgb = RGBColor(0, 128, 0)
        
        doc.add_paragraph()
    
    # Test Case 6: Affirmative & Negative Responses (CRITICAL)
    doc.add_heading('6. Affirmative & Negative Responses (CRITICAL)', 2)
    
    critical_note = doc.add_paragraph()
    critical_run = critical_note.add_run('⚠️ CRITICAL: These tests verify that duplicate responses have been fixed!')
    critical_run.font.bold = True
    critical_run.font.color.rgb = RGBColor(255, 0, 0)
    critical_run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    test_cases_affirm = [
        {
            'id': '6.1',
            'name': 'Yes After Insurance Query',
            'setup': ['Ask about insurance first: "tell me about insurance"', 'Then type: "yes"'],
            'prompt': 'yes',
            'expected': [
                '⚠️ SINGLE response (not duplicates)',
                'Shows detailed insurance plans',
                'From database if available'
            ],
            'pass_criteria': 'SINGLE response, shows insurance details',
            'critical': True
        },
        {
            'id': '6.2',
            'name': 'Yes After Doctor Query',
            'setup': ['Ask about doctors first: "I need a doctor"', 'Then type: "yes"'],
            'prompt': 'yes',
            'expected': [
                '⚠️ SINGLE response (not duplicates)',
                'Shows available doctors from database',
                'Offers to book appointment'
            ],
            'pass_criteria': 'SINGLE response, shows doctors',
            'critical': True
        },
        {
            'id': '6.3',
            'name': 'Generic Yes',
            'prompt': 'yes (without prior context)',
            'expected': [
                '⚠️ SINGLE response (not duplicates)',
                'Helpful menu of options',
                'Asks what user needs help with'
            ],
            'pass_criteria': 'SINGLE response, helpful guidance',
            'critical': True
        },
        {
            'id': '6.4',
            'name': 'Negative Response',
            'prompt': 'no',
            'expected': [
                'Acknowledges',
                'Asks what else they need help with'
            ],
            'pass_criteria': 'Single response, continues conversation'
        }
    ]
    
    for tc in test_cases_affirm:
        heading = f'Test Case {tc["id"]}: {tc["name"]}'
        if tc.get('critical'):
            heading += ' ⚠️ CRITICAL'
        doc.add_heading(heading, 3)
        
        # Setup if applicable
        if 'setup' in tc:
            p = doc.add_paragraph()
            p.add_run('Setup:').bold = True
            for step in tc['setup']:
                para = doc.add_paragraph(step, style='List Bullet 2')
                para.runs[0].font.color.rgb = RGBColor(255, 140, 0)
        
        # Prompt
        p = doc.add_paragraph()
        p.add_run('Prompt: ').bold = True
        p.add_run(f'"{tc["prompt"]}"').italic = True
        
        # Expected Response
        p = doc.add_paragraph()
        p.add_run('Expected Response:').bold = True
        for exp in tc['expected']:
            para = doc.add_paragraph(exp, style='List Bullet 2')
            if '⚠️' in exp or 'SINGLE' in exp:
                para.runs[0].font.color.rgb = RGBColor(255, 0, 0)
                para.runs[0].font.bold = True
        
        # Pass Criteria
        p = doc.add_paragraph()
        p.add_run('Pass Criteria: ').bold = True
        pass_run = p.add_run(f'✅ {tc["pass_criteria"]}')
        pass_run.font.color.rgb = RGBColor(0, 128, 0)
        pass_run.font.bold = True
        
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # Test Case 7: Specialized Healthcare Queries
    doc.add_heading('7. Specialized Healthcare Queries', 2)
    
    test_cases_specialized = [
        {
            'id': '7.1',
            'name': 'Lab Results',
            'prompt': 'I want to see my lab results',
            'expected': [
                'Offers to retrieve lab results',
                'Explains what results mean',
                'Guides on next steps'
            ],
            'pass_criteria': 'Helpful lab result information'
        },
        {
            'id': '7.2',
            'name': 'Billing Questions',
            'prompt': 'I have a billing question',
            'expected': [
                'Offers billing assistance',
                'Lists what it can help with (statements, charges, payment plans)',
                'Asks for specific billing question'
            ],
            'pass_criteria': 'Comprehensive billing help options'
        },
        {
            'id': '7.3',
            'name': 'Emergency',
            'prompt': 'This is an emergency',
            'expected': [
                'Immediate acknowledgment',
                'Directs to emergency services (911)',
                'Shows emergency contact numbers',
                'Offers urgent care options'
            ],
            'pass_criteria': 'Appropriate emergency handling'
        },
        {
            'id': '7.4',
            'name': 'Mental Health',
            'prompt': "I'm feeling anxious",
            'expected': [
                'Empathetic response',
                'Offers mental health assessments',
                'Provides crisis hotline information',
                'Offers to find mental health professionals'
            ],
            'pass_criteria': 'Empathetic, provides resources'
        },
        {
            'id': '7.5',
            'name': 'Wellness & Lifestyle',
            'prompt': 'I need diet recommendations',
            'expected': [
                'Offers personalized diet recommendations',
                'Asks about dietary preferences/restrictions',
                'Provides wellness guidance'
            ],
            'pass_criteria': 'Helpful wellness information'
        }
    ]
    
    for tc in test_cases_specialized:
        doc.add_heading(f'Test Case {tc["id"]}: {tc["name"]}', 3)
        
        p = doc.add_paragraph()
        p.add_run('Prompt: ').bold = True
        p.add_run(f'"{tc["prompt"]}"').italic = True
        
        p = doc.add_paragraph()
        p.add_run('Expected Response:').bold = True
        for exp in tc['expected']:
            doc.add_paragraph(exp, style='List Bullet 2')
        
        p = doc.add_paragraph()
        p.add_run('Pass Criteria: ').bold = True
        pass_run = p.add_run(f'✅ {tc["pass_criteria"]}')
        pass_run.font.color.rgb = RGBColor(0, 128, 0)
        
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # Critical Tests Section
    doc.add_heading('🎯 CRITICAL TESTS (MUST PASS)', 1)
    
    critical_intro = doc.add_paragraph()
    critical_intro_run = critical_intro.add_run('These tests MUST pass before the bot can be considered production-ready:')
    critical_intro_run.font.bold = True
    critical_intro_run.font.size = Pt(12)
    critical_intro_run.font.color.rgb = RGBColor(255, 0, 0)
    
    doc.add_paragraph()
    
    doc.add_heading('Priority 1: No Duplicates', 2)
    duplicates_tests = [
        'Test "yes" returns SINGLE response (not 10)',
        'Test "all plans" returns SINGLE response',
        'Test "suggest doctors" returns SINGLE response'
    ]
    for test in duplicates_tests:
        para = doc.add_paragraph(test, style='List Bullet')
        para.runs[0].font.bold = True
        para.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    
    doc.add_heading('Priority 2: Always Responds', 2)
    response_tests = [
        'Bot responds to every message (never silent)',
        'No "Sorry, I couldn\'t process" errors',
        'Shows helpful fallback for unrecognized queries'
    ]
    for test in response_tests:
        para = doc.add_paragraph(test, style='List Bullet')
        para.runs[0].font.bold = True
    
    doc.add_heading('Priority 3: Database Integration', 2)
    db_tests = [
        'Doctor searches retrieve from database',
        'Insurance plans show database data',
        'Responses reference specific database information'
    ]
    for test in db_tests:
        para = doc.add_paragraph(test, style='List Bullet')
        para.runs[0].font.bold = True
    
    doc.add_heading('Priority 4: Intelligent Responses', 2)
    intelligence_tests = [
        'Context-aware (remembers conversation)',
        '"Yes" understands what it refers to',
        'Uses AWS Bedrock for intelligent responses'
    ]
    for test in intelligence_tests:
        para = doc.add_paragraph(test, style='List Bullet')
        para.runs[0].font.bold = True
    
    doc.add_page_break()
    
    # Test Results Template
    doc.add_heading('📊 Test Results Template', 1)
    
    doc.add_paragraph('Use this template to record your test results:')
    doc.add_paragraph()
    
    # Add table
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    headers = ['Test Case', 'Prompt', 'Response Count', 'Quality', 'Pass/Fail', 'Notes']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add some sample rows
    test_rows = [
        ('1.1', 'Hello', '', '', '', ''),
        ('2.1', 'I am suffering from viral', '', '', '', ''),
        ('2.2', 'yes (after 2.1)', '', '', '', ''),
        ('3.1', 'suggest me some doctors', '', '', '', ''),
        ('4.1', 'I need help with insurance', '', '', '', ''),
        ('4.2', 'all plans', '', '', '', ''),
        ('6.1', 'yes (after insurance)', '', '', '', ''),
        ('6.2', 'yes (after doctor query)', '', '', '', ''),
        ('7.3', 'This is an emergency', '', '', '', ''),
    ]
    
    for test_row in test_rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(test_row):
            row_cells[i].text = value
    
    doc.add_paragraph()
    
    # Critical Checks Section
    doc.add_heading('CRITICAL CHECKS:', 2)
    
    critical_checks = [
        'All "yes" responses return SINGLE message (not 10 duplicates)',
        'Bot responds to every query (never silent)',
        'Responses are relevant and helpful',
        'Database data is shown when available'
    ]
    
    for check in critical_checks:
        para = doc.add_paragraph()
        para.add_run('☐ ').font.size = Pt(14)
        check_run = para.add_run(check)
        check_run.font.bold = True
        check_run.font.color.rgb = RGBColor(255, 0, 0)
    
    doc.add_page_break()
    
    # What to Look For Section
    doc.add_heading('🔍 What to Look For', 1)
    
    doc.add_heading('✅ PASS Indicators', 2)
    pass_indicators = [
        'Single responses for all queries (especially "yes")',
        'Bot responds within 2-3 seconds',
        'Responses are relevant and helpful',
        'Database data shown when available (doctors, insurance plans)',
        'No error messages visible to user',
        'Context maintained across conversation',
        'Professional, empathetic tone'
    ]
    for indicator in pass_indicators:
        para = doc.add_paragraph(indicator, style='List Bullet')
        para.runs[0].font.color.rgb = RGBColor(0, 128, 0)
    
    doc.add_heading('❌ FAIL Indicators', 2)
    fail_indicators = [
        'Multiple duplicate responses (especially for "yes")',
        '"Sorry, I couldn\'t process your message" errors',
        'Bot doesn\'t respond (silent)',
        'Irrelevant responses',
        'Generic responses when database data should be shown',
        'Error messages or stack traces visible',
        'Long delays (>10 seconds)'
    ]
    for indicator in fail_indicators:
        para = doc.add_paragraph(indicator, style='List Bullet')
        para.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    
    doc.add_page_break()
    
    # Success Criteria
    doc.add_heading('🎯 Success Criteria', 1)
    
    success_intro = doc.add_paragraph()
    success_intro_run = success_intro.add_run('Bot is production-ready if:')
    success_intro_run.font.bold = True
    success_intro_run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    success_criteria = [
        'All "yes" responses are SINGLE (not 10 duplicates)',
        'Bot responds to 100% of queries',
        'Responses are relevant and helpful',
        'Database integration works (shows doctors, plans from database)',
        'Context is maintained across conversation',
        'Error handling is graceful',
        'Response time is acceptable (<5 seconds)'
    ]
    
    for criterion in success_criteria:
        para = doc.add_paragraph()
        para.add_run('✅ ').font.size = Pt(14)
        criterion_run = para.add_run(criterion)
        criterion_run.font.bold = True
        criterion_run.font.color.rgb = RGBColor(0, 128, 0)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Final Notes
    doc.add_heading('🎉 Final Notes', 1)
    
    final_para = doc.add_paragraph('The bot has been enhanced with:')
    final_para.runs[0].font.bold = True
    
    enhancements = [
        'RAG (Retrieval Augmented Generation) for database integration',
        'AWS Bedrock LLM for super intelligent responses',
        'Duplicate prevention at multiple levels',
        'Context-aware conversation handling',
        'Comprehensive healthcare capabilities across all segments',
        'Robust error handling and fallback mechanisms'
    ]
    
    for enhancement in enhancements:
        para = doc.add_paragraph(enhancement, style='List Bullet')
        para.runs[0].font.color.rgb = RGBColor(0, 102, 204)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Contact/Support Info
    doc.add_heading('📞 Reporting Issues', 1)
    
    doc.add_paragraph('If you find any issues during testing, please report with:')
    
    reporting_items = [
        'Test case number',
        'Exact prompt used',
        'Expected vs actual response',
        'Screenshot if possible',
        'Timestamp of test',
        'Number of duplicate responses (if applicable)'
    ]
    
    for item in reporting_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Footer
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('Happy Testing! 🚀')
    footer_run.font.size = Pt(14)
    footer_run.font.bold = True
    footer_run.font.color.rgb = RGBColor(0, 102, 204)
    
    # Save document
    doc.save('CHATBOT_TESTING_GUIDE.docx')
    print("✅ Word document created successfully: CHATBOT_TESTING_GUIDE.docx")

if __name__ == '__main__':
    create_testing_document()

