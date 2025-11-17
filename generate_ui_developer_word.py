#!/usr/bin/env python3
"""
Generate Word Document for UI Developer Guide
Creates a comprehensive Word document with all UI integration information
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def add_heading_with_style(doc, text, level=1):
    """Add heading with custom styling"""
    heading = doc.add_heading(text, level=level)
    heading.style.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)
    return heading

def add_code_block(doc, code, language='javascript'):
    """Add code block with formatting"""
    p = doc.add_paragraph()
    p.style = 'No Spacing'
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0)
    # Add gray background
    p_format = p.paragraph_format
    p_format.left_indent = Inches(0.5)
    p_format.right_indent = Inches(0.5)
    p_format.space_before = Pt(6)
    p_format.space_after = Pt(6)
    return p

def add_table_of_contents(doc):
    """Add table of contents"""
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Overview',
        '2. API Endpoint Information',
        '3. Quick Start (5 Minutes)',
        '4. Frontend Integration Examples',
        '  4.1 React/TypeScript Examples',
        '  4.2 Vue.js Examples',
        '  4.3 Vanilla JavaScript Examples',
        '5. Environment Configuration',
        '6. Error Handling',
        '7. Testing the Integration',
        '8. Common Issues & Solutions',
        '9. Best Practices',
        '10. Quick Reference',
        '11. Running Backend Locally (Optional)'
    ]
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.style.font.size = Pt(11)

def create_ui_developer_word_document():
    """Create comprehensive UI Developer Word document"""
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title Page
    title = doc.add_heading('PRAN Chatbot - UI Developer Integration Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Complete Guide for Frontend Integration')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.paragraph_format
    subtitle_format.space_after = Pt(12)
    
    date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para_format = date_para.paragraph_format
    date_para_format.space_after = Pt(24)
    
    doc.add_page_break()
    
    # Table of Contents
    add_table_of_contents(doc)
    doc.add_page_break()
    
    # 1. Overview
    add_heading_with_style(doc, '1. Overview', 1)
    doc.add_paragraph(
        'This guide is specifically designed for UI developers who need to integrate '
        'the PRAN Healthcare Chatbot backend into their frontend applications. The chatbot '
        'provides a RESTful API that can be easily integrated into any modern frontend framework '
        'including React, Vue.js, Angular, or vanilla JavaScript.'
    )
    
    doc.add_paragraph(
        'As a UI developer, you have two options:',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Use the existing backend (if already deployed or running)',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        'Run the backend locally for testing (optional)',
        style='List Bullet 2'
    )
    
    # 2. API Endpoint Information
    add_heading_with_style(doc, '2. API Endpoint Information', 1)
    
    doc.add_paragraph('Primary Chat Endpoint:', style='Heading 3')
    add_code_block(doc, 'POST http://localhost:5001/rasa-webhook')
    
    doc.add_paragraph(
        'Note: Replace localhost:5001 with your production URL when deploying.'
    )
    
    doc.add_paragraph('Request Format:', style='Heading 3')
    add_code_block(doc, '''{
  "sender": "unique_user_id",
  "message": "Hello, I need a doctor"
}''')
    
    doc.add_paragraph('Response Format:', style='Heading 3')
    add_code_block(doc, '''[
  {
    "recipient_id": "unique_user_id",
    "text": "Hello! I'm your healthcare assistant. How can I help you today?"
  }
]''')
    
    doc.add_paragraph(
        'Important: The response is an array of message objects. Always handle it as an array.'
    )
    
    doc.add_paragraph('Health Check Endpoint:', style='Heading 3')
    add_code_block(doc, 'GET http://localhost:5001/health')
    
    add_code_block(doc, '''{
  "status": "healthy",
  "flask_wrapper": "running",
  "rasa_status": "connected",
  "mongodb_status": "connected"
}''')
    
    # 3. Quick Start
    add_heading_with_style(doc, '3. Quick Start (5 Minutes)', 1)
    
    doc.add_paragraph('Step 1: API Endpoint', style='Heading 3')
    add_code_block(doc, 'POST http://localhost:5001/rasa-webhook')
    
    doc.add_paragraph('Step 2: Send Request', style='Heading 3')
    add_code_block(doc, '''fetch('http://localhost:5001/rasa-webhook', {
  method: 'POST',
  headers: { 'Content-Type': 'application/json' },
  body: JSON.stringify({
    sender: 'user123',
    message: 'Hello'
  })
})''')
    
    doc.add_paragraph('Step 3: Handle Response', style='Heading 3')
    add_code_block(doc, '''const response = await fetch(...);
const messages = await response.json();
const botMessage = messages[0].text; // Extract bot response''')
    
    doc.add_paragraph('That\'s it! CORS is already enabled, so you can make requests from any frontend domain.')
    
    # 4. Frontend Integration Examples
    add_heading_with_style(doc, '4. Frontend Integration Examples', 1)
    
    # 4.1 React/TypeScript
    add_heading_with_style(doc, '4.1 React/TypeScript Examples', 2)
    
    doc.add_paragraph('Service Class:', style='Heading 3')
    add_code_block(doc, '''// ChatbotService.ts
const CHATBOT_API_URL = process.env.REACT_APP_CHATBOT_API_URL || 'http://localhost:5001';

export class ChatbotService {
  async sendMessage(userId: string, message: string): Promise<string[]> {
    try {
      const response = await fetch(`${CHATBOT_API_URL}/rasa-webhook`, {
        method: 'POST',
        headers: {
          'Content-Type': 'application/json',
        },
        body: JSON.stringify({
          sender: userId,
          message: message,
        }),
      });

      if (!response.ok) {
        throw new Error(`HTTP error! status: ${response.status}`);
      }

      const data = await response.json();
      return data.map((msg: any) => msg.text);
    } catch (error) {
      console.error('Error sending message:', error);
      throw error;
    }
  }

  async checkHealth(): Promise<boolean> {
    try {
      const response = await fetch(`${CHATBOT_API_URL}/health`);
      const data = await response.json();
      return data.status === 'healthy';
    } catch {
      return false;
    }
  }
}''')
    
    doc.add_paragraph('React Hook:', style='Heading 3')
    add_code_block(doc, '''// useChatbot.ts
import { useState } from 'react';

interface Message {
  text: string;
  sender: 'user' | 'bot';
}

export function useChatbot(userId: string = 'user123') {
  const [messages, setMessages] = useState<Message[]>([]);
  const [loading, setLoading] = useState(false);
  const [error, setError] = useState<string | null>(null);

  const sendMessage = async (message: string) => {
    if (!message.trim()) return;

    setLoading(true);
    setError(null);

    // Add user message immediately
    setMessages(prev => [...prev, { text: message, sender: 'user' }]);

    try {
      const response = await fetch(
        `${process.env.REACT_APP_CHATBOT_API_URL || 'http://localhost:5001'}/rasa-webhook`,
        {
          method: 'POST',
          headers: { 'Content-Type': 'application/json' },
          body: JSON.stringify({
            sender: userId,
            message: message,
          }),
        }
      );

      if (!response.ok) {
        throw new Error(`HTTP error! status: ${response.status}`);
      }

      const botResponses = await response.json();

      // Add bot responses
      botResponses.forEach((msg: any) => {
        setMessages(prev => [...prev, { text: msg.text, sender: 'bot' }]);
      });
    } catch (err) {
      const errorMessage = err instanceof Error ? err.message : 'Unknown error';
      setError(errorMessage);
      setMessages(prev => [...prev, { 
        text: 'Sorry, I encountered an error. Please try again.', 
        sender: 'bot' 
      }]);
    } finally {
      setLoading(false);
    }
  };

  return { messages, sendMessage, loading, error };
}''')
    
    doc.add_paragraph('React Component:', style='Heading 3')
    add_code_block(doc, '''// Chatbot.tsx
import React, { useState } from 'react';
import { useChatbot } from './useChatbot';

export function Chatbot() {
  const [input, setInput] = useState('');
  const { messages, sendMessage, loading } = useChatbot('user123');

  const handleSubmit = async (e: React.FormEvent) => {
    e.preventDefault();
    if (!input.trim() || loading) return;
    
    const message = input;
    setInput('');
    await sendMessage(message);
  };

  return (
    <div className="chatbot-container">
      <div className="messages">
        {messages.map((msg, idx) => (
          <div key={idx} className={`message ${msg.sender}`}>
            {msg.text}
          </div>
        ))}
        {loading && <div className="loading">Bot is typing...</div>}
      </div>
      
      <form onSubmit={handleSubmit}>
        <input
          type="text"
          value={input}
          onChange={(e) => setInput(e.target.value)}
          placeholder="Type your message..."
          disabled={loading}
        />
        <button type="submit" disabled={loading || !input.trim()}>
          Send
        </button>
      </form>
    </div>
  );
}''')
    
    # 4.2 Vue.js
    add_heading_with_style(doc, '4.2 Vue.js Examples', 2)
    add_code_block(doc, '''<template>
  <div class="chatbot">
    <div class="messages">
      <div 
        v-for="(msg, idx) in messages" 
        :key="idx" 
        :class="['message', msg.sender]"
      >
        {{ msg.text }}
      </div>
    </div>
    
    <form @submit.prevent="sendMessage">
      <input 
        v-model="input" 
        placeholder="Type your message..."
        :disabled="loading"
      />
      <button type="submit" :disabled="loading || !input.trim()">
        Send
      </button>
    </form>
  </div>
</template>

<script setup>
import { ref } from 'vue';

const messages = ref([]);
const input = ref('');
const loading = ref(false);
const userId = 'user123';
const API_URL = process.env.VUE_APP_CHATBOT_API_URL || 'http://localhost:5001';

const sendMessage = async () => {
  if (!input.value.trim() || loading.value) return;
  
  const userMessage = input.value;
  messages.value.push({ text: userMessage, sender: 'user' });
  input.value = '';
  loading.value = true;

  try {
    const response = await fetch(`${API_URL}/rasa-webhook`, {
      method: 'POST',
      headers: { 'Content-Type': 'application/json' },
      body: JSON.stringify({
        sender: userId,
        message: userMessage,
      }),
    });

    const botResponses = await response.json();
    botResponses.forEach(msg => {
      messages.value.push({ text: msg.text, sender: 'bot' });
    });
  } catch (error) {
    messages.value.push({ 
      text: 'Sorry, an error occurred. Please try again.', 
      sender: 'bot' 
    });
  } finally {
    loading.value = false;
  }
};
</script>''')
    
    # 4.3 Vanilla JavaScript
    add_heading_with_style(doc, '4.3 Vanilla JavaScript Examples', 2)
    add_code_block(doc, '''// chatbot.js
const CHATBOT_API_URL = 'http://localhost:5001';
const userId = 'user123';

async function sendMessage(message) {
  try {
    const response = await fetch(`${CHATBOT_API_URL}/rasa-webhook`, {
      method: 'POST',
      headers: {
        'Content-Type': 'application/json',
      },
      body: JSON.stringify({
        sender: userId,
        message: message,
      }),
    });

    const data = await response.json();
    return data.map(msg => msg.text);
  } catch (error) {
    console.error('Error:', error);
    throw error;
  }
}

// Usage
sendMessage('Hello')
  .then(responses => {
    responses.forEach(text => {
      console.log('Bot:', text);
      // Display in UI
    });
  })
  .catch(error => {
    console.error('Failed to send message:', error);
  });''')
    
    # 5. Environment Configuration
    add_heading_with_style(doc, '5. Environment Configuration', 1)
    
    doc.add_paragraph('For React (.env):', style='Heading 3')
    add_code_block(doc, 'REACT_APP_CHATBOT_API_URL=http://localhost:5001')
    
    doc.add_paragraph('For Vue (.env):', style='Heading 3')
    add_code_block(doc, 'VUE_APP_CHATBOT_API_URL=http://localhost:5001')
    
    doc.add_paragraph('For Production:', style='Heading 3')
    add_code_block(doc, 'REACT_APP_CHATBOT_API_URL=https://your-production-domain.com')
    
    # 6. Error Handling
    add_heading_with_style(doc, '6. Error Handling', 1)
    
    doc.add_paragraph('The API returns standard HTTP status codes:')
    doc.add_paragraph('200 OK - Success', style='List Bullet')
    doc.add_paragraph('400 Bad Request - Invalid request format', style='List Bullet')
    doc.add_paragraph('500 Internal Server Error - Server error', style='List Bullet')
    doc.add_paragraph('503 Service Unavailable - Backend services unavailable', style='List Bullet')
    
    doc.add_paragraph('Error Response Format:', style='Heading 3')
    add_code_block(doc, '''{
  "error": "Error message description"
}''')
    
    doc.add_paragraph('Example Error Handling:', style='Heading 3')
    add_code_block(doc, '''try {
  const response = await fetch(`${API_URL}/rasa-webhook`, {
    method: 'POST',
    headers: { 'Content-Type': 'application/json' },
    body: JSON.stringify({ sender: userId, message }),
  });

  if (!response.ok) {
    const error = await response.json();
    throw new Error(error.error || 'Failed to send message');
  }

  const messages = await response.json();
  return messages;
} catch (error) {
  console.error('Chatbot error:', error);
  // Show user-friendly error message
  throw error;
}''')
    
    # 7. Testing
    add_heading_with_style(doc, '7. Testing the Integration', 1)
    
    doc.add_paragraph('Test Health Endpoint:', style='Heading 3')
    add_code_block(doc, '''curl http://localhost:5001/health''')
    
    doc.add_paragraph('Test Chat Endpoint:', style='Heading 3')
    add_code_block(doc, '''curl -X POST http://localhost:5001/rasa-webhook \\
  -H "Content-Type: application/json" \\
  -d '{"sender": "test_user", "message": "Hello"}' ''')
    
    # 8. Common Issues
    add_heading_with_style(doc, '8. Common Issues & Solutions', 1)
    
    doc.add_paragraph('Issue 1: CORS Errors', style='Heading 3')
    doc.add_paragraph('Solution: CORS is enabled by default. If you see CORS errors, check that:')
    doc.add_paragraph('• Backend is running', style='List Bullet 2')
    doc.add_paragraph('• API URL is correct', style='List Bullet 2')
    
    doc.add_paragraph('Issue 2: Connection Refused', style='Heading 3')
    doc.add_paragraph('Solution:')
    doc.add_paragraph('• Backend is not running', style='List Bullet 2')
    doc.add_paragraph('• Wrong port (should be 5001 for Flask wrapper)', style='List Bullet 2')
    doc.add_paragraph('• Check backend health: curl http://localhost:5001/health', style='List Bullet 2')
    
    doc.add_paragraph('Issue 3: 503 Service Unavailable', style='Heading 3')
    doc.add_paragraph('Solution: Rasa server is not running. Backend needs both Flask wrapper AND Rasa server running.')
    
    doc.add_paragraph('Issue 4: Empty Response', style='Heading 3')
    doc.add_paragraph('Solution:')
    doc.add_paragraph('• Check response format - it\'s an array, not a single object', style='List Bullet 2')
    doc.add_paragraph('• Use response[0].text or map over the array', style='List Bullet 2')
    
    # 9. Best Practices
    add_heading_with_style(doc, '9. Best Practices', 1)
    
    doc.add_paragraph('1. Always check health before sending messages', style='List Bullet')
    add_code_block(doc, '''const isHealthy = await chatbot.checkHealth();
if (!isHealthy) {
  // Show error to user
}''')
    
    doc.add_paragraph('2. Handle loading states', style='List Bullet')
    doc.add_paragraph('• Show loading indicator while waiting for response', style='List Bullet 2')
    doc.add_paragraph('• Disable input during request', style='List Bullet 2')
    
    doc.add_paragraph('3. Error handling', style='List Bullet')
    doc.add_paragraph('• Always wrap API calls in try-catch', style='List Bullet 2')
    doc.add_paragraph('• Show user-friendly error messages', style='List Bullet 2')
    
    doc.add_paragraph('4. User ID management', style='List Bullet')
    doc.add_paragraph('• Use consistent user IDs (session ID, user ID, etc.)', style='List Bullet 2')
    doc.add_paragraph('• Store in localStorage or session storage', style='List Bullet 2')
    
    doc.add_paragraph('5. Message history', style='List Bullet')
    doc.add_paragraph('• Maintain conversation context', style='List Bullet 2')
    doc.add_paragraph('• Consider storing messages locally', style='List Bullet 2')
    
    # 10. Quick Reference
    add_heading_with_style(doc, '10. Quick Reference', 1)
    
    doc.add_paragraph('API Endpoints:', style='Heading 3')
    
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Endpoint'
    header_cells[1].text = 'Method'
    header_cells[2].text = 'Description'
    
    # Data rows
    data_rows = [
        ('/health', 'GET', 'Check backend status'),
        ('/rasa-webhook', 'POST', 'Send chat message'),
    ]
    
    for i, (endpoint, method, desc) in enumerate(data_rows, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = endpoint
        row_cells[1].text = method
        row_cells[2].text = desc
    
    doc.add_paragraph('Request Format:', style='Heading 3')
    add_code_block(doc, '''{
  "sender": "user_id",
  "message": "user message"
}''')
    
    doc.add_paragraph('Response Format:', style='Heading 3')
    add_code_block(doc, '''[
  {
    "recipient_id": "user_id",
    "text": "bot response"
  }
]''')
    
    # 11. Running Backend Locally
    add_heading_with_style(doc, '11. Running Backend Locally (Optional)', 1)
    
    doc.add_paragraph(
        'If you need to run the backend yourself for testing, follow these steps:'
    )
    
    doc.add_paragraph('Step 1: Clone Repository', style='Heading 3')
    add_code_block(doc, '''git clone https://github.com/PranDotAI1/pran_chatbot.git
cd pran_chatbot
git checkout new_pran_bot_aws''')
    
    doc.add_paragraph('Step 2: Run Automated Setup', style='Heading 3')
    add_code_block(doc, './setup_backend.sh')
    
    doc.add_paragraph('Step 3: Start Services (2 terminals needed)', style='Heading 3')
    doc.add_paragraph('Terminal 1 - Rasa Server:', style='Heading 4')
    add_code_block(doc, '''source venv/bin/activate
cd backend/app
rasa run --enable-api --cors "*" --port 5005''')
    
    doc.add_paragraph('Terminal 2 - Flask Wrapper:', style='Heading 4')
    add_code_block(doc, '''source venv/bin/activate
cd backend
python wrapper_server.py''')
    
    doc.add_paragraph('For detailed setup instructions, see DEVELOPER_SETUP_GUIDE.md in the repository.')
    
    # Footer
    doc.add_page_break()
    doc.add_paragraph('End of Document', style='Heading 1')
    doc.add_paragraph(
        'For additional help, refer to:',
        style='List Bullet'
    )
    doc.add_paragraph('• UI_DEVELOPER_GUIDE.md - Complete guide', style='List Bullet 2')
    doc.add_paragraph('• UI_QUICK_START.md - Quick reference', style='List Bullet 2')
    doc.add_paragraph('• TROUBLESHOOTING.md - Common issues', style='List Bullet 2')
    doc.add_paragraph('• Repository: https://github.com/PranDotAI1/pran_chatbot', style='List Bullet 2')
    
    return doc

def main():
    """Generate the Word document"""
    print("Generating UI Developer Word Document...")
    
    doc = create_ui_developer_word_document()
    
    filename = 'PRAN_Chatbot_UI_Developer_Guide.docx'
    doc.save(filename)
    
    print(f"✅ Word document created: {filename}")
    print(f"📄 File location: {filename}")

if __name__ == '__main__':
    main()

