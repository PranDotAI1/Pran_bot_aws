#!/usr/bin/env python3
"""
Generate comprehensive Word document report for New Pran Bot AWS
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
from datetime import datetime

def create_report():
    """Create comprehensive Word document report"""
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('New Pran Bot AWS - Comprehensive Technical Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('Production-Ready Healthcare Chatbot on AWS Infrastructure')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0].font
    subtitle_format.size = Pt(14)
    subtitle_format.italic = True
    
    # Date
    date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(10)
    
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        '1. Executive Summary',
        '2. System Overview',
        '3. AWS Services Integration',
        '4. Core Capabilities',
        '5. Medical Services',
        '6. Technical Architecture',
        '7. API Endpoints',
        '8. Example Responses',
        '9. Database Integration',
        '10. Security & Compliance',
        '11. Deployment Architecture',
        '12. Monitoring & Logging',
        '13. Development & Maintenance'
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', 1)
    doc.add_paragraph(
        'New Pran Bot AWS is a production-ready, enterprise-grade healthcare chatbot system '
        'deployed on Amazon Web Services (AWS) cloud infrastructure. The system leverages '
        'cutting-edge AI technologies including AWS Bedrock (Claude 3.5 Sonnet), AWS Comprehend Medical, '
        'and Rasa NLP framework to provide intelligent, context-aware healthcare assistance.'
    )
    
    doc.add_paragraph(
        'The chatbot serves as a comprehensive healthcare companion, capable of handling '
        'patient registration, appointment scheduling, symptom assessment, medication management, '
        'insurance verification, and much more. It integrates seamlessly with multiple databases '
        'including PostgreSQL (Aurora), MongoDB, and provides RESTful APIs for frontend integration.'
    )
    
    doc.add_paragraph(
        'Key highlights:',
        style='List Bullet'
    )
    highlights = [
        '100% AWS-native infrastructure with 25+ AWS services',
        'Production-ready codebase with no hardcoded credentials',
        'HIPAA-compliant medical data handling',
        'RAG (Retrieval-Augmented Generation) powered responses',
        'Multi-database support (PostgreSQL, MongoDB)',
        'Comprehensive error handling and logging',
        'Docker containerization for easy deployment',
        'Scalable microservices architecture'
    ]
    for highlight in highlights:
        doc.add_paragraph(highlight, style='List Bullet 2')
    
    doc.add_page_break()
    
    # 2. System Overview
    doc.add_heading('2. System Overview', 1)
    
    doc.add_heading('2.1 Architecture Components', 2)
    components = {
        'Frontend': 'React 18 application with TypeScript, providing modern user interface',
        'Backend': 'Rasa NLP engine with custom actions, handling natural language understanding',
        'API Gateway': 'Flask wrapper server providing RESTful API endpoints',
        'Actions Server': 'Python-based custom actions server for business logic',
        'Databases': 'PostgreSQL (Aurora) for relational data, MongoDB for document storage',
        'AI Services': 'AWS Bedrock for LLM, AWS Comprehend Medical for entity extraction'
    }
    
    for component, description in components.items():
        p = doc.add_paragraph()
        p.add_run(f'{component}: ').bold = True
        p.add_run(description)
    
    doc.add_heading('2.2 Technology Stack', 2)
    tech_stack = {
        'NLP Framework': 'Rasa 3.6.15',
        'AI/ML': 'AWS Bedrock (Claude 3.5 Sonnet), AWS Comprehend Medical',
        'Backend': 'Python 3.10, Flask, Rasa SDK',
        'Frontend': 'React 18, TypeScript, Vite',
        'Databases': 'PostgreSQL (Aurora RDS), MongoDB',
        'Containerization': 'Docker, Docker Compose',
        'Infrastructure': 'AWS ECS Fargate, Application Load Balancer',
        'Monitoring': 'AWS CloudWatch'
    }
    
    for tech, version in tech_stack.items():
        p = doc.add_paragraph()
        p.add_run(f'{tech}: ').bold = True
        p.add_run(version)
    
    doc.add_page_break()
    
    # 3. AWS Services Integration
    doc.add_heading('3. AWS Services Integration', 1)
    
    doc.add_paragraph(
        'The system leverages a comprehensive suite of AWS services for production-grade '
        'deployment and operation. All services are configured through environment variables '
        'with no hardcoded credentials, ensuring security and flexibility.'
    )
    
    aws_services = {
        'AWS Bedrock': {
            'Purpose': 'Large Language Model (LLM) for intelligent responses',
            'Model': 'Claude 3.5 Sonnet (anthropic.claude-3-5-sonnet-20241022-v2:0)',
            'Capabilities': [
                'Natural language understanding and generation',
                'Context-aware healthcare responses',
                'RAG (Retrieval-Augmented Generation) integration',
                'Multi-turn conversation handling',
                'Medical terminology comprehension'
            ]
        },
        'AWS Comprehend Medical': {
            'Purpose': 'Medical entity recognition and extraction',
            'Capabilities': [
                'Extract medications, conditions, anatomy, procedures',
                'Detect ICD-10 codes for medical conditions',
                'Detect RxNorm codes for medications',
                'Identify protected health information (PHI)',
                'Medical terminology understanding'
            ]
        },
        'AWS Comprehend': {
            'Purpose': 'Natural language processing',
            'Capabilities': [
                'Sentiment analysis',
                'Language detection',
                'Entity extraction',
                'Key phrase extraction'
            ]
        },
        'Amazon ECS Fargate': {
            'Purpose': 'Container orchestration',
            'Features': [
                'Serverless container hosting',
                'Auto-scaling capabilities',
                'High availability',
                'Easy deployment and management'
            ]
        },
        'Amazon RDS (Aurora PostgreSQL)': {
            'Purpose': 'Primary relational database',
            'Features': [
                'Managed PostgreSQL database',
                'Automatic backups',
                'High availability',
                'Multi-AZ deployment',
                'Connection pooling support'
            ]
        },
        'Amazon DocumentDB / MongoDB': {
            'Purpose': 'NoSQL document storage',
            'Features': [
                'MongoDB-compatible API',
                'Scalable document storage',
                'JSON document support',
                'High performance queries'
            ]
        },
        'Amazon ElastiCache (Redis)': {
            'Purpose': 'Caching and session management',
            'Features': [
                'In-memory data store',
                'Session caching',
                'Performance optimization',
                'Reduced database load'
            ]
        },
        'Application Load Balancer (ALB)': {
            'Purpose': 'Traffic distribution',
            'Features': [
                'Intelligent request routing',
                'Health checks',
                'SSL/TLS termination',
                'Auto-scaling support'
            ]
        },
        'Amazon S3': {
            'Purpose': 'Object storage',
            'Features': [
                'Static asset storage',
                'File uploads',
                'Backup storage',
                'CDN integration'
            ]
        },
        'AWS CloudWatch': {
            'Purpose': 'Monitoring and logging',
            'Features': [
                'Application logs',
                'Performance metrics',
                'Alarms and alerts',
                'Dashboard visualization'
            ]
        },
        'Amazon VPC': {
            'Purpose': 'Network isolation',
            'Features': [
                'Private subnets',
                'Security groups',
                'Network ACLs',
                'VPN connectivity'
            ]
        },
        'AWS Secrets Manager': {
            'Purpose': 'Credential management',
            'Features': [
                'Secure credential storage',
                'Automatic rotation',
                'IAM integration',
                'Encryption at rest'
            ]
        }
    }
    
    for service, details in aws_services.items():
        doc.add_heading(f'3.{list(aws_services.keys()).index(service) + 1} {service}', 2)
        if 'Purpose' in details:
            p = doc.add_paragraph()
            p.add_run('Purpose: ').bold = True
            p.add_run(details['Purpose'])
        if 'Model' in details:
            p = doc.add_paragraph()
            p.add_run('Model: ').bold = True
            p.add_run(details['Model'])
        if 'Features' in details:
            doc.add_paragraph('Features:', style='List Bullet')
            for feature in details['Features']:
                doc.add_paragraph(feature, style='List Bullet 2')
        if 'Capabilities' in details:
            doc.add_paragraph('Capabilities:', style='List Bullet')
            for capability in details['Capabilities']:
                doc.add_paragraph(capability, style='List Bullet 2')
    
    doc.add_page_break()
    
    # 4. Core Capabilities
    doc.add_heading('4. Core Capabilities', 1)
    
    doc.add_paragraph(
        'The chatbot provides comprehensive healthcare services through intelligent conversation. '
        'It uses RAG (Retrieval-Augmented Generation) to combine real-time database information '
        'with AI-generated responses for accurate, context-aware assistance.'
    )
    
    capabilities = {
        'Natural Language Understanding': [
            'Intent recognition for 40+ healthcare intents',
            'Entity extraction (dates, times, symptoms, medications)',
            'Context-aware conversation handling',
            'Multi-turn dialogue management',
            'Sentiment analysis for patient concerns'
        ],
        'Intelligent Response Generation': [
            'AWS Bedrock Claude 3.5 Sonnet integration',
            'RAG-powered responses with database context',
            'Medical entity recognition via Comprehend Medical',
            'Personalized recommendations',
            'Evidence-based medical guidance'
        ],
        'Database Integration': [
            'PostgreSQL/Aurora for relational data',
            'MongoDB for document storage',
            'Connection pooling for performance',
            'Automatic failover to backup databases',
            'Real-time data retrieval'
        ],
        'API Integration': [
            'RESTful API endpoints',
            'Webhook support for external systems',
            'MongoDB exploration APIs',
            'Health check endpoints',
            'CORS-enabled for frontend integration'
        ]
    }
    
    for capability, features in capabilities.items():
        doc.add_heading(f'4.{list(capabilities.keys()).index(capability) + 1} {capability}', 2)
        for feature in features:
            doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()
    
    # 5. Medical Services
    doc.add_heading('5. Medical Services', 1)
    
    medical_services = {
        'Medical Services': [
            'Symptom Assessment: Analyze symptoms, provide guidance, recommend appropriate care',
            'Emergency Assessment: Evaluate emergency situations, calculate HEART scores, assess urgency',
            'Mental Health: GAD-7 and PHQ-9 assessments, crisis detection, counselor recommendations',
            'Medication Management: Drug interaction checking, dosage calculations, medication schedules',
            'Health Education: Provide information about conditions, treatments, and wellness'
        ],
        'Patient Management': [
            'Patient Registration: Help register new patients with medical history and insurance',
            'Medical Records: Access and share medical records (HIPAA compliant)',
            'Patient History: View complete patient medical history',
            'Profile Management: Update patient profiles and information'
        ],
        'Appointments & Scheduling': [
            'Book Appointments: Schedule appointments with doctors',
            'Find Doctors: Match symptoms to appropriate medical specializations',
            'Reschedule/Cancel: Help manage existing appointments',
            'Appointment Reminders: Set up WhatsApp reminders',
            'Doctor Availability: Check doctor schedules and availability'
        ],
        'Insurance & Billing': [
            'Insurance Verification: Verify insurance coverage and benefits',
            'Insurance Plans: Show available insurance plans',
            'Insurance Suggestions: Recommend insurance based on needs',
            'Pre-authorization: Help with pre-authorization requests',
            'Cost Estimates: Provide service cost estimates',
            'Payment Plans: Explain payment plan options',
            'Billing Information: Access billing details and statements'
        ],
        'Wellness & Support': [
            'Diet Recommendations: Personalized diet plans',
            'Exercise Plans: Fitness and exercise recommendations',
            'Sleep Hygiene: Sleep quality tips and advice',
            'Clinical Guidelines: Evidence-based clinical recommendations'
        ],
        'Analytics & Administration': [
            'Disease Trends: Analyze health trends and patterns',
            'Feedback Collection: Gather patient feedback',
            'Health Predictions: Predictive health analytics',
            'Health Recommendations: Personalized health advice'
        ],
        'Hospital Services': [
            'Hospital Locations: Find nearby hospitals and clinics',
            'Hospital Policies: Explain hospital policies and procedures',
            'Country Services: Location-specific healthcare services'
        ]
    }
    
    for category, services in medical_services.items():
        doc.add_heading(f'5.{list(medical_services.keys()).index(category) + 1} {category}', 2)
        for service in services:
            doc.add_paragraph(service, style='List Bullet')
    
    doc.add_page_break()
    
    # 6. Technical Architecture
    doc.add_heading('6. Technical Architecture', 1)
    
    doc.add_heading('6.1 System Architecture', 2)
    doc.add_paragraph(
        'The system follows a microservices architecture with clear separation of concerns:'
    )
    
    architecture_components = [
        ('Frontend Layer', 'React application serving the user interface, communicates with Flask API gateway'),
        ('API Gateway', 'Flask wrapper server handling HTTP requests, routing to Rasa backend'),
        ('NLP Engine', 'Rasa server processing natural language, understanding intents and entities'),
        ('Actions Server', 'Custom Python actions executing business logic, database queries, AWS service calls'),
        ('Database Layer', 'PostgreSQL for structured data, MongoDB for documents, Redis for caching'),
        ('AI Services', 'AWS Bedrock for LLM, Comprehend Medical for entity extraction'),
        ('Infrastructure', 'ECS Fargate containers, Load Balancer, VPC networking')
    ]
    
    for component, description in architecture_components:
        p = doc.add_paragraph()
        p.add_run(f'{component}: ').bold = True
        p.add_run(description)
    
    doc.add_heading('6.2 Data Flow', 2)
    doc.add_paragraph('1. User sends message via Frontend')
    doc.add_paragraph('2. Frontend calls Flask API Gateway (/rasa-webhook)')
    doc.add_paragraph('3. Flask forwards request to Rasa backend')
    doc.add_paragraph('4. Rasa processes intent and entities')
    doc.add_paragraph('5. Rasa calls Actions Server for custom logic')
    doc.add_paragraph('6. Actions Server queries databases and AWS services')
    doc.add_paragraph('7. AWS Bedrock generates intelligent response')
    doc.add_paragraph('8. Response flows back through layers to user')
    
    doc.add_heading('6.3 RAG (Retrieval-Augmented Generation) System', 2)
    doc.add_paragraph(
        'The system implements RAG to provide accurate, context-aware responses:'
    )
    rag_steps = [
        'User query is analyzed for intent and entities',
        'Relevant context is retrieved from databases (doctors, appointments, medical records)',
        'Medical entities are extracted using AWS Comprehend Medical',
        'Retrieved context is combined with user query',
        'AWS Bedrock generates response using both context and general knowledge',
        'Response is personalized based on retrieved information'
    ]
    for step in rag_steps:
        doc.add_paragraph(step, style='List Bullet')
    
    doc.add_page_break()
    
    # 7. API Endpoints
    doc.add_heading('7. API Endpoints', 1)
    
    doc.add_heading('7.1 Flask Wrapper Server Endpoints', 2)
    
    endpoints = {
        'GET /health': {
            'Description': 'Health check endpoint',
            'Response': 'Returns status of Flask wrapper, Rasa backend, and MongoDB connection',
            'Example': '{"status": "healthy", "flask_wrapper": "running", "rasa_status": "connected", "mongodb_status": "connected"}'
        },
        'POST /rasa-webhook': {
            'Description': 'Main chat endpoint - forwards requests to Rasa',
            'Request Body': '{"sender": "user_id", "message": "user message", "metadata": {}}',
            'Response': 'Array of bot responses from Rasa'
        },
        'GET /mongodb/test': {
            'Description': 'Test MongoDB connection',
            'Response': 'Returns connection status and list of databases'
        },
        'GET /mongodb/explore': {
            'Description': 'Explore MongoDB structure',
            'Response': 'Returns databases, collections, document counts, and sample documents'
        }
    }
    
    for endpoint, details in endpoints.items():
        p = doc.add_paragraph()
        p.add_run(endpoint).bold = True
        for key, value in details.items():
            doc.add_paragraph(f'{key}: {value}', style='List Bullet 2')
    
    doc.add_heading('7.2 Rasa Backend Endpoints', 2)
    rasa_endpoints = {
        'POST /webhooks/rest/webhook': 'Main webhook for chat messages',
        'GET /status': 'Rasa server status',
        'POST /model/parse': 'Parse user message for intent and entities'
    }
    
    for endpoint, description in rasa_endpoints.items():
        p = doc.add_paragraph()
        p.add_run(endpoint).bold = True
        p.add_run(f': {description}')
    
    doc.add_page_break()
    
    # 8. Example Responses
    doc.add_heading('8. Example Responses', 1)
    
    doc.add_heading('8.1 Symptom Assessment', 2)
    examples = [
        ('Cold and Cough', 
         'I understand you\'re experiencing cold and cough symptoms. For immediate relief, rest, stay hydrated, and consider over-the-counter cold medications. Gargle with warm salt water for cough relief. If symptoms persist for more than 10 days, you have a high fever (over 101°F), difficulty breathing, or severe headache, please see a doctor. I can help you book an appointment with a doctor or find an ENT specialist. Would you like me to help you find a doctor or book an appointment?'),
        
        ('Fever',
         'I understand you have a fever. Normal body temperature is 98.6°F (37°C), fever is 100.4°F (38°C) or higher, and high fever is 103°F (39.4°C) or higher. Please see a doctor if your fever is over 103°F, lasts more than 3 days, or is accompanied by severe symptoms like rash, difficulty breathing, or confusion. For infants under 3 months, any fever requires immediate medical attention. I can help you book an appointment with a doctor or find available urgent care. Would you like me to help you schedule a consultation?'),
        
        ('Headache',
         'I understand you\'re experiencing a headache. Common causes include tension headaches, migraines, sinus issues, dehydration, or stress. If you have a sudden severe headache (worst of your life), headache with fever/stiff neck/confusion, headache after head injury, or vision changes, seek immediate care. I can help you book an appointment with a neurologist or general physician. Would you like me to find a doctor or schedule a consultation?')
    ]
    
    for symptom, response in examples:
        p = doc.add_paragraph()
        p.add_run(f'Symptom: {symptom}').bold = True
        doc.add_paragraph(f'Response: {response}')
        doc.add_paragraph('')
    
    doc.add_heading('8.2 Insurance Information', 2)
    insurance_response = """Available Insurance Plans:

1. Basic Health Plan
   Monthly Premium: $150
   Deductible: $1000
   Coverage: 80%
   Features: Primary care, Emergency visits, Basic prescriptions

2. Premium Health Plan
   Monthly Premium: $300
   Deductible: $500
   Coverage: 90%
   Features: All basic features, Specialist visits, Mental health, Dental & Vision

3. Family Health Plan
   Monthly Premium: $450
   Deductible: $750
   Coverage: 85%
   Features: All premium features, Family coverage, Maternity care, Pediatric care

Would you like more details about any specific plan, or would you like personalized insurance recommendations based on your needs?"""
    
    doc.add_paragraph(insurance_response)
    
    doc.add_heading('8.3 Doctor Finding', 2)
    doc.add_paragraph(
        'When users ask to find a doctor, the system:'
    )
    doctor_finding_steps = [
        'Extracts specialty from user message (cardiologist, neurologist, etc.)',
        'Queries database for matching doctors',
        'Falls back to API if database unavailable',
        'Presents list of available doctors with specialties, departments, contact info',
        'Offers to book appointment with selected doctor'
    ]
    for step in doctor_finding_steps:
        doc.add_paragraph(step, style='List Bullet')
    
    doc.add_page_break()
    
    # 9. Database Integration
    doc.add_heading('9. Database Integration', 1)
    
    doc.add_heading('9.1 PostgreSQL (Aurora)', 2)
    postgres_tables = [
        'doctors: Doctor profiles, specialties, contact information',
        'patients: Patient records, demographics, medical history',
        'appointments: Appointment scheduling, status, notes',
        'medical_records: Diagnosis, treatment, prescriptions',
        'insurance_plans: Insurance plan details, coverage, premiums'
    ]
    doc.add_paragraph('Key Tables:')
    for table in postgres_tables:
        doc.add_paragraph(table, style='List Bullet')
    
    doc.add_heading('9.2 MongoDB', 2)
    doc.add_paragraph(
        'MongoDB is used for document storage and exploration. The system can connect to '
        'multiple MongoDB databases and explore their structure through API endpoints.'
    )
    doc.add_paragraph('MongoDB Features:')
    mongodb_features = [
        'Multi-database support',
        'Collection exploration',
        'Document sampling',
        'Field analysis',
        'Connection testing'
    ]
    for feature in mongodb_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('9.3 Database Connection Management', 2)
    doc.add_paragraph(
        'The system implements robust database connection management:'
    )
    db_features = [
        'Connection pooling for PostgreSQL',
        'Automatic failover to backup databases',
        'Connection timeout handling',
        'Error recovery mechanisms',
        'Environment variable configuration'
    ]
    for feature in db_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()
    
    # 10. Security & Compliance
    doc.add_heading('10. Security & Compliance', 1)
    
    doc.add_heading('10.1 Security Measures', 2)
    security_measures = [
        'All credentials stored in environment variables (no hardcoded values)',
        'AWS Secrets Manager for sensitive data',
        'HTTPS/TLS for all external communications',
        'CORS properly configured for frontend access',
        'Input validation on all API endpoints',
        'SQL injection prevention through parameterized queries',
        'Network isolation through VPC and security groups',
        'IAM roles with least privilege principle'
    ]
    for measure in security_measures:
        doc.add_paragraph(measure, style='List Bullet')
    
    doc.add_heading('10.2 HIPAA Compliance', 2)
    doc.add_paragraph(
        'The system is designed with HIPAA compliance in mind:'
    )
    hipaa_features = [
        'Protected Health Information (PHI) handling',
        'Audit logging for medical record access',
        'Secure data transmission',
        'Access control mechanisms',
        'Data encryption at rest and in transit'
    ]
    for feature in hipaa_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('10.3 Code Security', 2)
    code_security = [
        'No hardcoded passwords or API keys',
        'Environment variable templates provided',
        'Comprehensive .gitignore to prevent credential leaks',
        'Production-ready error messages (no sensitive info)',
        'Structured logging without sensitive data'
    ]
    for item in code_security:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 11. Deployment Architecture
    doc.add_heading('11. Deployment Architecture', 1)
    
    doc.add_heading('11.1 Container Architecture', 2)
    doc.add_paragraph('The system is containerized using Docker:')
    containers = [
        'rasa: Rasa NLP backend server (Port 5005)',
        'rasa-actions: Custom actions server (Port 5055)',
        'flask-wrapper: API gateway server (Port 5001)'
    ]
    for container in containers:
        doc.add_paragraph(container, style='List Bullet')
    
    doc.add_heading('11.2 AWS Deployment', 2)
    aws_deployment = [
        'ECS Fargate: Serverless container hosting',
        'Application Load Balancer: Traffic distribution and SSL termination',
        'Auto Scaling: Automatic scaling based on demand',
        'Multi-AZ: High availability across availability zones',
        'CloudWatch: Monitoring and alerting'
    ]
    for item in aws_deployment:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('11.3 Infrastructure as Code', 2)
    doc.add_paragraph(
        'Deployment configurations are managed through:'
    )
    iac_items = [
        'Terraform: Infrastructure provisioning',
        'Docker Compose: Local development and testing',
        'Environment templates: Configuration management',
        'Deployment scripts: Automated deployment processes'
    ]
    for item in iac_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 12. Monitoring & Logging
    doc.add_heading('12. Monitoring & Logging', 1)
    
    doc.add_heading('12.1 Logging', 2)
    logging_features = [
        'Structured logging using Python logging module',
        'Log levels: INFO, WARNING, ERROR, DEBUG',
        'CloudWatch Logs integration',
        'Request/response logging',
        'Error tracking and stack traces'
    ]
    for feature in logging_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('12.2 Monitoring', 2)
    monitoring_features = [
        'Health check endpoints for all services',
        'CloudWatch Metrics for performance tracking',
        'Database connection monitoring',
        'API response time tracking',
        'Error rate monitoring',
        'Custom CloudWatch dashboards'
    ]
    for feature in monitoring_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('12.3 Alerts', 2)
    doc.add_paragraph('Alerts can be configured for:')
    alerts = [
        'Service downtime',
        'High error rates',
        'Database connection failures',
        'AWS service unavailability',
        'Performance degradation'
    ]
    for alert in alerts:
        doc.add_paragraph(alert, style='List Bullet')
    
    doc.add_page_break()
    
    # 13. Development & Maintenance
    doc.add_heading('13. Development & Maintenance', 1)
    
    doc.add_heading('13.1 Code Quality', 2)
    quality_features = [
        'Production-ready codebase (no emojis, no hardcoding)',
        'Comprehensive error handling',
        'Type hints where applicable',
        'Code documentation and docstrings',
        'Modular architecture for maintainability'
    ]
    for feature in quality_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('13.2 Testing', 2)
    testing_info = [
        'Rasa model testing: rasa test',
        'Python syntax validation',
        'API endpoint testing',
        'Database connection testing',
        'Integration testing capabilities'
    ]
    for item in testing_info:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('13.3 Maintenance', 2)
    maintenance_items = [
        'Regular dependency updates',
        'Security patch management',
        'Database backup procedures',
        'Log rotation and archival',
        'Performance optimization'
    ]
    for item in maintenance_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # Conclusion
    doc.add_page_break()
    doc.add_heading('Conclusion', 1)
    doc.add_paragraph(
        'New Pran Bot AWS represents a comprehensive, production-ready healthcare chatbot solution '
        'leveraging the full power of AWS cloud infrastructure and cutting-edge AI technologies. '
        'The system provides intelligent, context-aware healthcare assistance while maintaining '
        'high standards for security, compliance, and scalability.'
    )
    doc.add_paragraph(
        'With its RAG-powered responses, multi-database integration, and extensive AWS service '
        'utilization, the chatbot serves as a complete healthcare companion capable of '
        'handling a wide range of patient needs from appointment scheduling to symptom assessment '
        'to insurance management.'
    )
    doc.add_paragraph(
        'The production-ready codebase, comprehensive documentation, and robust architecture '
        'make this system suitable for enterprise deployment and continuous improvement.'
    )
    
    # Save document
    output_file = 'New_Pran_Bot_AWS_Comprehensive_Report.docx'
    doc.save(output_file)
    print(f'Report generated successfully: {output_file}')
    return output_file

if __name__ == '__main__':
    try:
        create_report()
    except ImportError:
        print("Error: python-docx library not installed.")
        print("Please install it using: pip install python-docx")
    except Exception as e:
        print(f"Error generating report: {e}")

